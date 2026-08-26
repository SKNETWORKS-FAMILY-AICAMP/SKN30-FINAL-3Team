#!/usr/bin/env python3
"""Create the two temporary submission DOCX files from actual PoC results.

The previous cohort documents are used as retained visual references. Their
body contents are replaced, while section geometry, styles, headers, footers,
theme relationships, and embedded branding remain source-derived.
"""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4FB"
LIGHT_GRAY = "E7E6E6"
MID_GRAY = "B4C6E7"
WHITE = "FFFFFF"
TEXT = "222222"

FONT_KO = "맑은 고딕"
FONT_EN = "Arial"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--poc-root", type=Path, required=True)
    parser.add_argument("--training-template", type=Path, required=True)
    parser.add_argument("--model-template", type=Path, required=True)
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Output directory. Defaults to <PoC root>/documents.",
    )
    parser.add_argument(
        "--version-suffix",
        default="",
        help="Optional suffix appended before .docx, for example _v2.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Allow replacing an existing output. Disabled by default.",
    )
    parser.add_argument("--submission-date", default="2026. 08. 20.")
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def load_comparison(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as file_obj:
        return list(csv.DictReader(file_obj))


def set_run_font(run: Any, name: str = FONT_KO) -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), name)


def configure_style(style: Any, size: float, bold: bool = False, color: str = TEXT) -> None:
    style.font.name = FONT_KO
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    rfonts.set(qn("w:eastAsia"), FONT_KO)


def configure_document(doc: DocumentObject, title: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = "F2 Field Proposal Review Risk Model synthetic PoC"
    doc.core_properties.author = "SK 네트웍스 Family AI 30기 3팀"
    doc.core_properties.comments = (
        "Synthetic proxy-label PoC. Not production or real-user performance."
    )

    configure_style(doc.styles["Normal"], 10.5, False, TEXT)
    normal_format = doc.styles["Normal"].paragraph_format
    normal_format.space_after = Pt(5)
    normal_format.line_spacing = 1.25

    if "Title" in doc.styles:
        configure_style(doc.styles["Title"], 24, True, NAVY)
    if "Subtitle" in doc.styles:
        configure_style(doc.styles["Subtitle"], 12, False, BLUE)
    if "Heading 1" in doc.styles:
        configure_style(doc.styles["Heading 1"], 15, True, NAVY)
        doc.styles["Heading 1"].paragraph_format.space_before = Pt(14)
        doc.styles["Heading 1"].paragraph_format.space_after = Pt(7)
        doc.styles["Heading 1"].paragraph_format.keep_with_next = True
    if "Heading 2" in doc.styles:
        configure_style(doc.styles["Heading 2"], 12, True, BLUE)
        doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
        doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)
        doc.styles["Heading 2"].paragraph_format.keep_with_next = True
    if "Heading 3" in doc.styles:
        configure_style(doc.styles["Heading 3"], 10.5, True, NAVY)


def clear_body(doc: DocumentObject) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_row_cant_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_fixed(table: Any, widths_cm: Sequence[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    if widths_cm:
        for row in table.rows:
            for index, width in enumerate(widths_cm):
                if index < len(row.cells):
                    row.cells[index].width = Cm(width)


def set_table_borders(table: Any, color: str = "B7C5D8", size: str = "6") -> None:
    """Apply explicit grid borders without depending on a localized table style."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table(table: Any, widths_cm: Sequence[float] | None = None) -> None:
    set_table_fixed(table, widths_cm)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_table_row_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                if row_index < len(table.rows) - 1:
                    # All generated tables fit on one page. Keeping each row
                    # with the following row avoids orphaned last/first rows.
                    paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    set_run_font(run)
                    run.font.size = Pt(8.8)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def add_table(
    doc: DocumentObject,
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
    widths_cm: Sequence[float] | None = None,
) -> Any:
    materialized = [["" if value is None else str(value) for value in row] for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_values in materialized:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    format_table(table, widths_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc: DocumentObject, text: str, level: int = 1) -> Any:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_body(doc: DocumentObject, text: str, bold_prefix: str | None = None) -> Any:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.font.bold = True
        set_run_font(prefix)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc: DocumentObject, text: str) -> Any:
    # The supplied Korean templates do not expose Word's localized bullet style
    # through the English name, so use an explicit bullet for deterministic output.
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(f"• {text}")
    set_run_font(run)
    return paragraph


def add_numbered(doc: DocumentObject, text: str) -> Any:
    style = "List Number" if "List Number" in doc.styles else "Normal"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc: DocumentObject, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=BLUE, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}  ")
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    set_run_font(title_run)
    body_run = paragraph.add_run(text)
    set_run_font(body_run)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: DocumentObject, image_path: Path, caption: str, width_inches: float = 5.7) -> None:
    if not image_path.is_file():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.font.size = Pt(9)
    caption_run.font.italic = True
    caption_run.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(caption_run)


def add_cover(
    doc: DocumentObject,
    title: str,
    subtitle: str,
    submission_date: str,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(22)
    cohort = doc.add_paragraph()
    cohort.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cohort.add_run("SK 네트웍스 Family AI 30기 : 3팀")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(run)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(26)
    title_paragraph.paragraph_format.space_after = Pt(8)
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    set_run_font(title_run)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_font(subtitle_run)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["산출물 단계", "데이터 전처리"],
            ["제출 일자", submission_date],
            ["깃허브 경로", "https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN30-FINAL-3Team"],
            ["작성 팀원", "3팀 전체"],
            ["기능명", "Field Proposal Review Risk Model"],
        ],
        [4.2, 12.0],
    )
    add_callout(
        doc,
        "제출 범위",
        "기존 프로젝트의 합성 상담 시나리오에서 규칙 기반 대리 라벨을 파생한 소규모 ML PoC이다. 실제 사용자 데이터나 실서비스 성능을 의미하지 않는다.",
    )
    doc.add_page_break()


def metric_rows(comparison: list[dict[str, str]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in comparison:
        rows.append(
            [
                row["Model"],
                f"{float(row['Train Accuracy']):.4f}",
                f"{float(row['Test Accuracy']):.4f}",
                f"{float(row['Precision']):.4f}",
                f"{float(row['Recall']):.4f}",
                f"{float(row['F1']):.4f}",
            ]
        )
    return rows


def format_pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def create_training_result_document(
    template: Path,
    output: Path,
    poc_root: Path,
    metrics: dict[str, Any],
    metadata: dict[str, Any],
    inventory: dict[str, Any],
    eda: dict[str, Any],
    comparison: list[dict[str, str]],
    submission_date: str,
) -> None:
    doc = Document(str(template))
    clear_body(doc)
    configure_document(doc, "머신러닝 / 딥러닝 학습 결과서")
    add_cover(
        doc,
        "머신러닝 / 딥러닝 학습 결과서",
        "F2 Field Proposal Review Risk Model - 강사 피드백 반영 v2",
        submission_date,
    )

    add_heading(doc, "1. 개요", 1)
    add_heading(doc, "1.1 실험 목적", 2)
    add_body(
        doc,
        "F2의 sLLM이 만든 장부 필드 제안에 대해 데이터 생성·전처리·분류 Pipeline을 재현할 수 있는지 확인하였다. 현재 needs_review는 실제 사용자 행동이 아니라 입력 Feature로 만든 대리 Target이므로, 현 단계의 평가 질문은 ‘규칙으로 정의한 위험도를 단순 모델이 재현할 수 있는가’이다. 실제 사용자의 수정·거절 행동 예측은 독립 라벨을 확보한 후 수행할 차기 실험으로 분리한다.",
    )
    add_table(
        doc,
        ["구분", "정의"],
        [
            ["Task", "Binary Classification"],
            ["Target 0", "LOW_RISK - field_type별 proxy risk score 하위 8건"],
            ["Target 1", "NEEDS_REVIEW - field_type별 proxy risk score 상위 7건"],
            ["Positive class", "needs_review=1"],
            ["PoC 선정 순서", "Test Recall → Test F1 → |Train-Test Accuracy 차이| → 단순성"],
        ],
        [4.0, 12.2],
    )
    add_callout(
        doc,
        "중요 해석",
        "Target 생성에 사용한 confidence·충돌·부정·언급 수·파싱 성공 여부를 다시 모델 입력으로 사용하였다. 따라서 현재 지표는 사용자 행동 예측력이 아니라 대리 라벨 생성 규칙의 재현 정도를 보여준다.",
    )

    add_heading(doc, "1.2 작업 범위", 2)
    for item in (
        "합성 상담 원천의 중복 제거와 150건 필드 제안 데이터 파생",
        "간단한 EDA와 80:20 층화 분할",
        "Dummy, Logistic Regression, Random Forest 비교",
        "혼동행렬·Feature Importance·모델 비교 결과 생성",
        "선정 Pipeline 전체 joblib 저장과 재로드 추론 확인",
    ):
        add_bullet(doc, item)
    add_body(doc, "딥러닝·Transformer·XGBoost·SHAP·외부 API·서비스 연동은 이번 범위에서 제외하였다.")

    doc.add_page_break()
    add_heading(doc, "2. 데이터와 전처리", 1)
    add_heading(doc, "2.1 원천 데이터", 2)
    add_body(
        doc,
        f"세 원천 파일에는 총 {inventory['raw_row_count']:,}행이 있으며, 파일 간 완전 중복 {inventory['duplicate_row_count']:,}건을 제거한 뒤 {inventory['deduplicated_row_count']:,}개의 고유 합성 상담을 후보 풀로 사용하였다.",
    )
    source_rows = []
    for source in inventory["sources"]:
        source_rows.append(
            [
                Path(source["path"]).name,
                source["raw_rows"],
                source["retained_rows_after_cross_file_deduplication"],
                ", ".join(f"{key} {value}" for key, value in source["raw_label_distribution"].items()),
            ]
        )
    add_table(
        doc,
        ["원천 파일", "원본 행", "중복 제거 후", "라벨 분포"],
        source_rows,
        [6.6, 2.0, 2.3, 5.5],
    )
    add_callout(
        doc,
        "50건을 제거한 이유",
        "50행 파일 f2_sell_request_scenarios.privacy_safe.v0.2.jsonl 전체가 200행 파일 f2_sllm_data_small.jsonl의 매도의뢰 50행과 완전히 같다. 50건 모두 scenario_id, source_group_id, transcript SHA-256과 전체 JSON 내용이 각각 일치해 별도 관측치가 아니므로 첫 출현만 남겼다.",
    )
    add_body(
        doc,
        "이 처리는 이상치나 내용 오류를 이유로 행을 버린 것이 아니다. 동일 상담이 후보 선택에서 두 번 가중되거나 Train/Test 양쪽에 나뉘는 중복 누수를 막기 위한 파일 간 중복 제거이며, 50행의 정보는 먼저 읽은 200행 파일에 그대로 보존되어 있다. 결과적으로 원천 1,250행은 고유 상담 1,200건으로 정리되었다. 현재 생성기는 세 키 중 하나만 겹쳐도 뒤 행을 제외하므로, 향후 일부 키만 충돌하고 내용이 다른 경우에는 자동 제거하지 않고 충돌 목록으로 격리해 확인하도록 개선한다.",
    )
    add_body(
        doc,
        "모든 원천 레코드는 contains_real_personal_data=false로 표시된 합성 문장이다. 다만 사람 검수된 운영 데이터가 아니므로 실제 상담 분포를 대표한다고 가정하지 않았다.",
    )

    add_heading(doc, "2.2 필드 제안 단위 데이터 생성", 2)
    for index, step in enumerate(
        (
            "상담 문장에서 금액·거래유형·동·호·평형·날짜·명도 관련 후보 근거를 정규식과 키워드로 탐지한다.",
            "각 field_type마다 서로 다른 상담 15건을 선택하고 같은 원천 그룹과 같은 transcript hash를 중복 사용하지 않는다.",
            "근거 길이, 후보 수, 충돌·부정 여부, 파싱 성공 여부와 모사 confidence를 산출한다.",
            "proxy risk score를 계산하고 field_type별 상위 7건을 NEEDS_REVIEW로 지정하여 총 70건의 positive class를 구성한다.",
        ),
        start=1,
    ):
        add_numbered(doc, f"{index}. {step}")

    add_body(
        doc,
        "proxy risk score = 2×(1-confidence) + 1.5×has_conflict + 1.2×has_negation + 0.35×(mention_count-1) + 1.4×(1-parse_success) + 0.75×no_fill_context + N(0, 0.35)",
        bold_prefix="proxy risk score",
    )

    add_table(
        doc,
        ["Feature", "형식", "정의"],
        [
            ["field_type", "범주형", "10개 장부 필드 유형"],
            ["confidence", "0~1 실수", "파싱·충돌·부정·후보 수와 seed 42로 모사한 신뢰도"],
            ["evidence_length", "5~80 정수", "근거 문장 길이"],
            ["mention_count", "1~4 정수", "해당 필드 후보 언급 수"],
            ["has_conflict", "0/1", "복수 값 또는 정정·변경·대조 표현 존재"],
            ["has_negation", "0/1", "부정·거부·보류 표현 존재"],
            ["parse_success", "0/1", "필드별 정규화기가 구체 값을 생성했는지"],
            ["needs_review", "0/1", "규칙 기반 대리 Target"],
        ],
        [4.2, 3.0, 9.2],
    )

    add_heading(doc, "2.3 Target 생성 변수와 Feature 중첩", 2)
    add_table(
        doc,
        ["구분", "현재 설계", "영향"],
        [
            [
                "직접 중첩",
                "confidence, mention_count, has_conflict, has_negation, parse_success",
                "7개 모델 입력 중 5개가 Target 점수식에 직접 사용됨",
            ],
            [
                "간접 중첩",
                "confidence 자체도 parse_success·conflict·negation·mention_count로 생성",
                "같은 신호가 confidence와 점수식에 반복 반영됨",
            ],
            [
                "층화 규칙",
                "field_type별 점수 상위 7건을 class 1로 지정",
                "field_type도 라벨 배정 과정에 관여함",
            ],
            [
                "직접 미사용",
                "evidence_length",
                "점수식에는 없지만 합성 원천에서 다른 신호와 연관될 수 있음",
            ],
        ],
        [3.2, 7.2, 6.0],
    )
    add_body(
        doc,
        "원천 그룹과 transcript의 Train/Test 교차가 0건이라는 검사는 행 중복 누수는 막지만, 위와 같은 Target 정의의 순환성은 해소하지 못한다. 이는 미래 정보가 입력에 섞인 전형적인 시점 누수라기보다 ‘라벨 생성 규칙과 입력 변수의 중첩’이며, 모델이 독립적인 정답 대신 알려진 규칙을 근사하도록 만든다.",
    )
    add_callout(
        doc,
        "현재 성능이 의미하는 범위",
        "Accuracy 0.8000은 실제 사용자의 검토 행동을 맞힌 비율이 아니다. 규칙 기반 합성 Target을 고정된 30건 Test Set에서 재현한 비율이며, Pipeline 동작 확인용 기준선으로만 사용한다.",
    )

    add_heading(doc, "2.4 데이터 품질과 분포", 2)
    add_table(
        doc,
        ["검사 항목", "결과"],
        [
            ["전체 행", eda["row_count"]],
            ["필드별 행", "10개 field_type 각각 15건"],
            ["LOW_RISK", eda["target_distribution"]["0"]],
            ["NEEDS_REVIEW", eda["target_distribution"]["1"]],
            ["Positive 비율", format_pct(float(eda["target_positive_ratio"]))],
            ["결측치", sum(int(value) for value in eda["missing_values"].values())],
            ["원천 그룹/문장 중복", "선택 데이터 0건"],
        ],
        [6.2, 10.2],
    )
    add_callout(doc, "재현성", f"random_state={metadata['random_state']}를 고정하고 데이터·생성기·원천 파일 SHA-256을 함께 기록하였다.")

    doc.add_page_break()
    add_heading(doc, "3. 실험 설계", 1)
    add_heading(doc, "3.1 전처리와 분할", 2)
    add_body(
        doc,
        f"needs_review 기준 층화 분할을 사용해 Train {metrics['split']['train_rows']}건, Test {metrics['split']['test_rows']}건으로 나누고 Test에 10개 field_type이 모두 포함되는지 별도 검증하였다. 같은 source_group_id와 transcript hash의 Train/Test 교차는 모두 {metrics['split']['source_group_overlap']}건이다.",
    )
    add_table(
        doc,
        ["구분", "처리"],
        [
            ["field_type", "OneHotEncoder(handle_unknown='ignore')"],
            ["Logistic 숫자 Feature", "SimpleImputer(median) + StandardScaler"],
            ["Random Forest 숫자 Feature", "SimpleImputer(median), 원 스케일 유지"],
            ["분할", "Train 80% / Test 20%, random_state=42"],
        ],
        [5.5, 10.9],
    )

    add_heading(doc, "3.2 모델 후보", 2)
    add_table(
        doc,
        ["모델", "주요 설정", "역할"],
        [
            ["Dummy", "most_frequent", "다수 클래스 기준선"],
            ["Logistic Regression", "max_iter=1000, random_state=42", "단순하고 영향 방향을 설명하기 쉬운 선형 분류"],
            ["Random Forest", "n_estimators=200, max_depth=6, random_state=42", "비선형 상호작용 비교 후보"],
        ],
        [4.4, 6.1, 5.9],
    )
    add_heading(doc, "3.3 평가 지표", 2)
    add_table(
        doc,
        ["지표", "해석", "사용 목적"],
        [
            ["Accuracy", "전체 예측 중 정답 비율", "전반적 성능 확인"],
            ["Precision", "검토 필요 예측 중 실제 대리 positive 비율", "불필요 검토 부담 확인"],
            ["Recall", "대리 검토 필요 항목 중 탐지 비율", "위험 항목 누락 최소화, 1순위"],
            ["F1", "Precision과 Recall의 조화 평균", "균형 성능, 2순위"],
        ],
        [3.2, 8.0, 5.2],
    )

    add_heading(doc, "4. 실험 결과", 1)
    add_heading(doc, "4.1 모델 비교", 2)
    add_table(
        doc,
        ["Model", "Train Acc.", "Test Acc.", "Precision", "Recall", "F1"],
        metric_rows(comparison),
        [4.0, 2.5, 2.5, 2.4, 2.2, 2.2],
    )
    add_body(doc, "모든 수치는 고정된 30건 합성 Test Set에서 실제 실행한 결과이며, 일반화된 실서비스 성능이 아니다.")
    add_callout(
        doc,
        "정확도 해석",
        "Random Forest가 더 높은 값은 Train Accuracy(0.9750 대 0.7917)이다. Test Accuracy는 두 모델 모두 0.8000이고 Precision 0.9000, Recall 0.6429, F1 0.7500, 혼동행렬(TN 15·FP 1·FN 5·TP 9)도 동일하다. 높은 학습 정확도만으로 일반화 성능이 더 좋다고 판단하지 않았다.",
    )

    add_heading(doc, "4.2 혼동행렬", 2)
    add_figure(
        doc,
        poc_root / "reports/confusion_matrix_logistic_regression.png",
        "그림 1. Logistic Regression confusion matrix",
        4.8,
    )
    add_figure(
        doc,
        poc_root / "reports/confusion_matrix_random_forest.png",
        "그림 2. Random Forest confusion matrix",
        4.8,
    )

    add_heading(doc, "4.3 Random Forest Feature Importance", 2)
    add_figure(
        doc,
        poc_root / "reports/random_forest_feature_importance.png",
        "그림 3. 변환된 Feature 기준 Random Forest 중요도",
        5.8,
    )
    importance_rows = []
    for item in metrics.get("random_forest_feature_importance", [])[:8]:
        importance_rows.append([item["feature"], f"{float(item['importance']):.4f}"])
    if importance_rows:
        add_table(doc, ["상위 Feature", "Importance"], importance_rows, [11.0, 5.2])
    add_body(
        doc,
        "Feature Importance는 이 규칙 기반 합성 데이터 안에서의 분할 기여도이며 인과관계나 실제 업무 중요도를 의미하지 않는다. Target 점수식에 직접 사용된 다섯 입력(confidence·has_conflict·mention_count·parse_success·has_negation)의 중요도 합이 약 0.7100이므로, 중요도 순위 역시 사용자 행동의 원인이 아니라 규칙 재현 구조를 주로 반영한다.",
    )

    add_heading(doc, "5. PoC 기준 모델 선정(임시)", 1)
    final_name = metrics["final_model"]["name"]
    final_result = metrics["models"][final_name]
    add_callout(doc, "PoC 기준 모델", final_name)
    add_body(
        doc,
        "사전에 정한 우선순위는 Test Recall, Test F1, |Train-Test Accuracy 차이|, 모델 단순성 순이다. 두 후보의 Test Recall과 F1뿐 아니라 Test Accuracy와 혼동행렬까지 같았으므로, 학습 정확도가 아니라 일반화 차이와 단순성을 다음 판단 근거로 사용하였다.",
    )
    logistic_result = metrics["models"]["Logistic Regression"]
    forest_result = metrics["models"]["Random Forest"]
    add_table(
        doc,
        ["판단 항목", "Logistic Regression", "Random Forest", "판단"],
        [
            ["Test Accuracy", f"{logistic_result['test_accuracy']:.4f}", f"{forest_result['test_accuracy']:.4f}", "동일"],
            ["Test Recall", f"{logistic_result['recall']:.4f}", f"{forest_result['recall']:.4f}", "동일"],
            ["Test F1", f"{logistic_result['f1']:.4f}", f"{forest_result['f1']:.4f}", "동일"],
            ["Train Accuracy", f"{logistic_result['train_accuracy']:.4f}", f"{forest_result['train_accuracy']:.4f}", "RF가 높음"],
            ["|Train-Test Acc. 차이|", f"{abs(logistic_result['train_test_accuracy_gap']):.4f}", f"{abs(forest_result['train_test_accuracy_gap']):.4f}", "LR이 작음"],
            ["구조", "선형·단순·계수 확인 가능", "비선형·상대적으로 복잡", "LR 우선"],
        ],
        [4.6, 4.3, 4.0, 3.3],
    )
    add_body(
        doc,
        "따라서 Logistic Regression은 현재 규칙 재현 PoC의 저장 artifact로 임시 선택하였다. Test 30건과 순환적 대리 Target만으로 Random Forest보다 실제 성능이 우수하다고 결론 내린 것은 아니며, 독립적인 사용자 라벨 데이터에서는 모델 선정부터 다시 수행한다.",
    )
    add_table(
        doc,
        ["선정 모델 Test 지표", "값"],
        [
            ["Accuracy", f"{final_result['test_accuracy']:.4f}"],
            ["Precision", f"{final_result['precision']:.4f}"],
            ["Recall", f"{final_result['recall']:.4f}"],
            ["F1", f"{final_result['f1']:.4f}"],
            ["|Train-Test Accuracy 차이|", f"{abs(final_result['train_test_accuracy_gap']):.4f}"],
        ],
        [9.0, 7.2],
    )
    add_body(doc, "선정된 전처리·모델 Pipeline 전체는 field_proposal_review_risk_model.joblib로 저장했으며 재로드 후 알려진 입력과 미지 field_type 입력 모두에서 0/1 예측을 확인하였다.")

    add_heading(doc, "6. 한계 및 개선 방향", 1)
    add_heading(doc, "6.1 현재 결과의 한계", 2)
    for limitation in (
        "실제 사용자 데이터가 아닌 합성 상담 시나리오를 사용하였다.",
        "needs_review와 confidence가 실제 관측값이 아니며, Target 생성 변수와 모델 입력이 중첩된다.",
        "데이터 150건, Test 30건으로 표본이 매우 작다.",
        "실제 F2 sLLM 오류 및 사용자 검토 분포와 다를 수 있다.",
        "현재 지표는 규칙 재현 성능이며 실제 사용자 행동 예측 또는 실서비스 성능이 아니다.",
        "현재 모델은 제출용 PoC이며 F2 서비스의 자동 반영 판단에 사용할 수 없다.",
        "딥러닝 모델은 이번 범위에서 학습하지 않았다.",
    ):
        add_bullet(doc, limitation)

    add_heading(doc, "6.2 차기 실험 계획(미실행)", 2)
    for index, step in enumerate(
        (
            "F2 검토 화면에서 제안별 사용자 최종 행동을 수집한다. 변경 없이 승인한 항목은 0, 수정 또는 거절한 항목은 1로 두고, 검토 중단·미확인 항목은 unknown으로 분리해 학습에서 제외한다.",
            "needs_review를 confidence·충돌·파싱 규칙으로 생성하지 않는다. 이 값들은 사용자 검토 전에 관측되는 후보 Feature로만 사용하며 Target과 독립시킨다.",
            "현재 proxy risk 수식을 명시적 Rule Baseline으로 구현하고, Logistic Regression·Random Forest 및 규칙 변수 제거 ablation을 같은 평가셋에서 비교한다.",
            "운영 사용자 라벨을 바로 확보할 수 없으면 규칙과 Feature 값을 보지 않은 복수 검수자가 독립 판정하고, 판정 불가와 불일치 해소 절차 및 평가자 일치도를 함께 기록한다.",
            "동일 사용자·원천 상담·파생 제안이 분할을 넘지 않도록 group split을 적용하고, 개발 비교에는 반복 group validation을, 최종 판정에는 시간 순서 Holdout을 사용해 변동 범위도 보고한다.",
            "Accuracy만으로 선택하지 않고 needs_review Recall, Precision, F1, 혼동행렬과 사용자 검토 부담을 함께 보고 임계값을 결정한다.",
        ),
        start=1,
    ):
        add_numbered(doc, f"{index}. {step}")
    add_callout(
        doc,
        "v2 반영 범위",
        "이번 v2는 기존 실험을 재해석하고 다음 검증 설계를 명시한 문서 개선본이다. 실제 사용자 라벨 수집·재학습·재평가는 아직 수행하지 않았다.",
    )

    add_heading(doc, "7. 주요 산출물", 1)
    add_table(
        doc,
        ["분류", "파일"],
        [
            ["데이터", "ml/field_proposal_reliability/data/synthetic_field_proposals.csv"],
            ["Notebook", "ml/field_proposal_reliability/notebooks/field_proposal_review_risk_poc_output.ipynb"],
            ["모델", "ml/field_proposal_reliability/artifacts/field_proposal_review_risk_model.joblib"],
            ["메타데이터", "ml/field_proposal_reliability/artifacts/model_metadata.json"],
            ["지표", "ml/field_proposal_reliability/reports/model_comparison.csv, metrics.json"],
            ["요약", "ml/field_proposal_reliability/reports/experiment_summary.md"],
        ],
        [4.0, 12.2],
    )
    add_heading(doc, "8. 변경 이력", 1)
    add_table(
        doc,
        ["변경일", "변경자", "변경내용", "버전"],
        [
            ["2026. 08. 20", "3팀 전체", "합성 대리 라벨 기반 제출용 ML PoC 결과서 작성", "v1.0"],
            [submission_date.rstrip("."), "3팀 전체", "중복 제거 근거, Target-Feature 중첩, 모델 선정 근거와 차기 실험 계획 보완", "v2.0"],
        ],
        [3.1, 3.2, 7.8, 2.1],
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def create_model_document(
    template: Path,
    output: Path,
    poc_root: Path,
    metrics: dict[str, Any],
    metadata: dict[str, Any],
    inventory: dict[str, Any],
    comparison: list[dict[str, str]],
    submission_date: str,
) -> None:
    doc = Document(str(template))
    clear_body(doc)
    configure_document(doc, "학습한 ML / DL 모델")
    add_cover(
        doc,
        "학습한 ML / DL 모델",
        "Field Proposal Review Risk Model - 강사 피드백 반영 v2",
        submission_date,
    )

    final_name = metrics["final_model"]["name"]
    final_result = metrics["models"][final_name]
    add_heading(doc, "1. 모델 개요", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["모델명", final_name],
            ["선정 상태", "규칙 재현 PoC의 임시 기준 모델 - 운영 최종 모델 아님"],
            ["기능명", "Field Proposal Review Risk Model"],
            ["현재 목적", "규칙 기반 proxy Target을 재현하는 학습·평가 Pipeline 확인"],
            ["차기 목적", "독립적인 사용자 수정·거절 라벨을 이용한 검토 필요 행동 예측"],
            ["문제 유형", "Binary Classification"],
            ["Positive class", "needs_review=1 (NEEDS_REVIEW)"],
            ["데이터", "기존 합성 상담에서 파생한 규칙 기반 대리 라벨 150건"],
        ],
        [4.2, 12.0],
    )

    add_heading(doc, "2. 모델 구조", 1)
    add_body(doc, "합성 상담 → 필드 후보 탐지 → 정형 Feature 생성 및 proxy Target 생성 → 80:20 층화 분할 → Rule 재현 모델 비교 → PoC 기준 모델 임시 선택 → Pipeline joblib 저장")

    add_heading(doc, "3. 입력 / 출력 정의", 1)
    add_table(
        doc,
        ["구분", "항목", "설명"],
        [
            ["입력", "field_type", "10개 장부 필드 유형"],
            ["입력", "confidence", "0~1 모사 신뢰도"],
            ["입력", "evidence_length", "근거 문장 길이"],
            ["입력", "mention_count", "후보 언급 수"],
            ["입력", "has_conflict / has_negation", "충돌·부정 표현 여부"],
            ["입력", "parse_success", "구조화 파싱 성공 여부"],
            ["출력", "0", "LOW_RISK"],
            ["출력", "1", "NEEDS_REVIEW"],
        ],
        [2.6, 5.2, 8.4],
    )

    add_heading(doc, "4. Target 정의", 1)
    add_body(doc, "needs_review는 실제 사용자 피드백이 아니라 아래 proxy risk score의 field_type별 상위 7건으로 생성하였다. 각 field_type 15건 중 7건이 class 1, 8건이 class 0이어서 전체 분포는 70:80이다.")
    add_body(
        doc,
        "proxy risk score = 2×(1-confidence) + 1.5×has_conflict + 1.2×has_negation + 0.35×(mention_count-1) + 1.4×(1-parse_success) + 0.75×no_fill_context + N(0, 0.35)",
        bold_prefix="proxy risk score",
    )
    add_heading(doc, "4.1 Target-Feature 중첩과 해석 제한", 2)
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["직접 재사용", "confidence, mention_count, has_conflict, has_negation, parse_success가 Target 생성과 모델 입력에 모두 사용됨"],
            ["간접 재사용", "confidence도 parse_success·conflict·negation·mention_count로 모사되어 같은 신호가 반복 반영됨"],
            ["층화 관여", "field_type별 점수 순위로 class를 배정하므로 field_type도 라벨 생성 과정에 관여함"],
            ["직접 미사용", "evidence_length는 점수식에 직접 사용되지 않음"],
        ],
        [4.0, 12.2],
    )
    add_callout(
        doc,
        "라벨 한계",
        "현재 모델은 독립적인 사용자 정답을 예측한 것이 아니라 입력 변수로 만든 규칙을 근사하였다. 따라서 Accuracy 0.8000은 사용자 행동 예측 성능이나 실서비스 성능으로 일반화할 수 없다.",
    )

    add_heading(doc, "5. 전처리와 사용 프레임워크", 1)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["범주형", "OneHotEncoder(handle_unknown='ignore')"],
            ["숫자형", "median imputation; Logistic Regression은 StandardScaler 적용"],
            ["Train/Test", f"{metrics['split']['train_rows']} / {metrics['split']['test_rows']}"],
            ["라이브러리", "pandas, NumPy, scikit-learn, Matplotlib, joblib"],
            ["난수 시드", str(metadata["random_state"])],
        ],
        [5.0, 11.2],
    )

    add_heading(doc, "6. 학습 모델과 평가", 1)
    add_table(
        doc,
        ["Model", "Train Acc.", "Test Acc.", "Precision", "Recall", "F1"],
        metric_rows(comparison),
        [4.0, 2.5, 2.5, 2.4, 2.2, 2.2],
    )
    add_body(doc, "검토 필요 항목 누락을 줄이기 위해 Test Recall을 우선하였다. 두 후보 모델은 Test Accuracy 0.8000, Precision 0.9000, Recall 0.6429, F1 0.7500과 혼동행렬(TN 15·FP 1·FN 5·TP 9)이 모두 같았다.")
    add_callout(
        doc,
        "Random Forest 정확도에 대한 설명",
        "Random Forest가 더 높은 것은 Train Accuracy(0.9750)이며 Test Accuracy는 Logistic Regression과 같은 0.8000이다. 학습 정확도 상승은 일반화 성능 향상의 근거가 아니고, Random Forest의 |Train-Test Accuracy 차이| 0.1750은 Logistic Regression의 0.0083보다 크다.",
    )
    add_figure(
        doc,
        poc_root
        / (
            "reports/confusion_matrix_logistic_regression.png"
            if final_name == "Logistic Regression"
            else "reports/confusion_matrix_random_forest.png"
        ),
        f"그림 1. 최종 모델({final_name}) confusion matrix",
        4.6,
    )

    add_heading(doc, "7. PoC 기준 최종 모델(임시)", 1)
    add_callout(doc, "PoC 기준 모델", final_name)
    logistic_result = metrics["models"]["Logistic Regression"]
    forest_result = metrics["models"]["Random Forest"]
    add_table(
        doc,
        ["판단 항목", "Logistic Regression", "Random Forest", "선택 해석"],
        [
            ["Test Accuracy", f"{logistic_result['test_accuracy']:.4f}", f"{forest_result['test_accuracy']:.4f}", "동일"],
            ["Test Recall", f"{logistic_result['recall']:.4f}", f"{forest_result['recall']:.4f}", "동일"],
            ["Test F1", f"{logistic_result['f1']:.4f}", f"{forest_result['f1']:.4f}", "동일"],
            ["Train Accuracy", f"{logistic_result['train_accuracy']:.4f}", f"{forest_result['train_accuracy']:.4f}", "RF가 높음"],
            ["|Train-Test Acc. 차이|", f"{abs(logistic_result['train_test_accuracy_gap']):.4f}", f"{abs(forest_result['train_test_accuracy_gap']):.4f}", "LR이 작음"],
            ["복잡도·설명", "선형·계수 확인 가능", "비선형·상대적으로 복잡", "LR 우선"],
        ],
        [4.4, 4.2, 4.0, 3.6],
    )
    add_body(
        doc,
        "Test 성능이 동률이므로 더 높은 Train Accuracy를 선택 근거로 사용하지 않았다. |Train-Test Accuracy 차이|가 작고 구조가 단순한 Logistic Regression을 현재 규칙 재현 PoC의 저장 artifact로 임시 선택하였다. 이 선택은 실제 사용자 라벨 데이터에서 다시 검증해야 한다.",
    )
    add_table(
        doc,
        ["Test 지표", "값"],
        [
            ["Accuracy", f"{final_result['test_accuracy']:.4f}"],
            ["Precision", f"{final_result['precision']:.4f}"],
            ["Recall", f"{final_result['recall']:.4f}"],
            ["F1", f"{final_result['f1']:.4f}"],
        ],
        [8.2, 8.0],
    )
    add_body(doc, "전체 전처리 Pipeline과 분류기를 field_proposal_review_risk_model.joblib로 저장하였다. 입력은 정의된 7개 Feature를 가진 pandas DataFrame이며 출력은 0 또는 1이다.")

    add_heading(doc, "8. 적용 전략", 1)
    for item in (
        "현재 모델은 제출용 오프라인 artifact로만 보관한다.",
        "F2 서비스 API나 사용자 승인·저장 흐름에는 연결하지 않는다.",
        "차기 Target은 변경 없는 승인=0, 사용자 수정 또는 거절=1로 정의하고 검토 중단·미확인은 unknown으로 분리한다.",
        "confidence·충돌·부정·언급 수·파싱 성공 여부는 Target 생성 규칙이 아니라 검토 전 관측 Feature로만 사용한다.",
        "현재 proxy risk 수식을 Rule Baseline으로 분리하고 ML 모델 및 규칙 변수 제거 ablation과 비교한다.",
        "동일 사용자·원천 상담·파생 제안의 group 누수를 막고 반복 group validation과 시간 순서 Holdout에서 변동 범위를 함께 평가한다.",
        "Accuracy만이 아니라 Recall·Precision·F1·혼동행렬과 사용자 검토 부담으로 모델과 임계값을 고른다.",
    ):
        add_bullet(doc, item)
    add_callout(doc, "계획 상태", "사용자 라벨 수집·재학습·재평가는 이번 v2 문서 범위에서 아직 실행하지 않았다.")

    add_heading(doc, "9. 한계 및 향후 개선", 1)
    for item in (
        "50행 원천 파일은 200행 원천 파일과 세 중복 키가 모두 일치하는 완전 부분집합이어서 중복 누수 방지를 위해 제거하였다.",
        "학습 데이터는 150건, Test Set은 30건으로 작다.",
        "confidence와 needs_review가 실제 관측값이 아니고 Target 생성 변수와 입력 Feature가 중첩된다.",
        "실제 F2 sLLM 오류 분포 및 사용자의 검토 행동과 다를 수 있다.",
        "본 결과는 규칙 재현용 소규모 합성 Test Set 결과이며 사용자 행동 예측 또는 실서비스 성능을 의미하지 않는다.",
        "실제 사용자 수정·거절 라벨과 독립적인 시간 기반 평가셋에서 모델 선정부터 재검증해야 한다.",
        "딥러닝은 이번 PoC 범위에서 수행하지 않았다.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "10. 모델 Artifact", 1)
    add_table(
        doc,
        ["파일", "설명"],
        [
            ["field_proposal_review_risk_model.joblib", "최종 전처리·분류 Pipeline"],
            ["model_metadata.json", "Feature·데이터 해시·선정 사유·라이브러리 버전"],
            ["metrics.json", "모델별 실제 평가 지표와 혼동행렬"],
        ],
        [9.0, 7.2],
    )

    add_heading(doc, "11. 변경 이력", 1)
    add_table(
        doc,
        ["변경일", "변경자", "변경내용", "버전"],
        [
            ["2026. 08. 20", "3팀 전체", "규칙 기반 대리 라벨 ML PoC 모델 설명서 작성", "v1.0"],
            [submission_date.rstrip("."), "3팀 전체", "Target-Feature 중첩과 Logistic Regression 임시 선정 근거 보완", "v2.0"],
        ],
        [3.1, 3.2, 7.8, 2.1],
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    args = parse_args()
    poc_root = args.poc_root.resolve()
    metrics = load_json(poc_root / "reports/metrics.json")
    metadata = load_json(poc_root / "artifacts/model_metadata.json")
    inventory = load_json(poc_root / "data/source_inventory.json")
    eda = load_json(poc_root / "reports/eda_summary.json")
    comparison = load_comparison(poc_root / "reports/model_comparison.csv")

    if any(character in args.version_suffix for character in ("/", "\\")):
        raise ValueError("--version-suffix must not contain path separators")
    documents_dir = (args.output_dir or poc_root / "documents").resolve()
    detailed_output = documents_dir / (
        f"[데이터 전처리] 머신러닝_딥러닝 학습 결과서_30기_3팀{args.version_suffix}.docx"
    )
    model_output = documents_dir / (
        f"[데이터 전처리] 학습한 ML_DL 모델_30기_3팀{args.version_suffix}.docx"
    )

    output_pairs = (
        (args.training_template.resolve(), detailed_output),
        (args.model_template.resolve(), model_output),
    )
    for template, output in output_pairs:
        if template == output:
            raise ValueError(f"Refusing to overwrite the template: {template}")
        if output.exists() and not args.overwrite:
            raise FileExistsError(
                f"Output already exists: {output}. Pass --overwrite to replace it."
            )

    create_training_result_document(
        args.training_template.resolve(),
        detailed_output,
        poc_root,
        metrics,
        metadata,
        inventory,
        eda,
        comparison,
        args.submission_date,
    )
    create_model_document(
        args.model_template.resolve(),
        model_output,
        poc_root,
        metrics,
        metadata,
        inventory,
        comparison,
        args.submission_date,
    )
    print(detailed_output)
    print(model_output)


if __name__ == "__main__":
    main()
